"""Multi-format document parser with OCR and Table Extraction support."""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any

logger = logging.getLogger(__name__)

class DocumentChunk:
    def __init__(self, text: str, metadata: Dict[str, Any]):
        self.text = text
        self.metadata = metadata

class DocumentParser:
    """Parses PDF, DOCX, and TXT files, extracting text and tables."""
    
    def __init__(self, use_ocr: bool = True):
        self.use_ocr = use_ocr
        self._ocr_warning_printed = False
        
    def parse_file(self, file_path: str) -> List[DocumentChunk]:
        """Parse a file and return a list of chunks (e.g. pages or sections)."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext in [".docx", ".doc"]:
            return self._parse_docx(path)
        elif ext == ".txt":
            return self._parse_txt(path)
        else:
            logger.warning(f"Unsupported file extension: {ext}")
            return []

    def _parse_pdf(self, path: Path) -> List[DocumentChunk]:
        chunks = []
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = "\n".join([" | ".join([str(cell).replace('\n', ' ') if cell else "" for cell in row]) for row in table])
                        text += f"\n\n[Extracted Table]\n{table_text}\n"
                    
                    # If text is still empty and OCR is enabled, try OCR
                    if not text.strip() and self.use_ocr:
                        text = self._ocr_pdf_page(path, i)
                        
                    if text.strip():
                        chunks.append(DocumentChunk(
                            text=text,
                            metadata={"source": path.name, "page": i + 1}
                        ))
        except ImportError:
            logger.error("pdfplumber not installed. Cannot parse PDF.")
        except Exception as e:
            logger.error(f"Error parsing PDF {path}: {e}")
            
        return chunks

    def _ocr_pdf_page(self, path: Path, page_num: int) -> str:
        if not self.use_ocr:
            return ""
            
        try:
            import pytesseract
            import fitz  # PyMuPDF
            from PIL import Image
            import io
            
            # Use PyMuPDF to convert PDF page to image (removes Poppler dependency)
            doc = fitz.open(path)
            page = doc.load_page(page_num)
            pix = page.get_pixmap(dpi=150)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            
            return pytesseract.image_to_string(img)
            
        except Exception as e:
            if not getattr(self, '_ocr_warning_printed', False):
                logger.warning(f"OCR missing or failed (Tesseract may not be installed on Windows). Disabling OCR to prevent spam. Error: {e}")
                self._ocr_warning_printed = True
            self.use_ocr = False
        return ""

    def _parse_docx(self, path: Path) -> List[DocumentChunk]:
        chunks = []
        try:
            import docx
            doc = docx.Document(path)
            
            text_blocks = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_blocks.append(para.text)
                    
            # Extract tables
            for table in doc.tables:
                table_text = "\n[Extracted Table]\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip().replace("\n", " ") for cell in row.cells])
                    table_text += row_text + "\n"
                text_blocks.append(table_text)
                
            full_text = "\n\n".join(text_blocks)
            
            if full_text.strip():
                chunks.append(DocumentChunk(
                    text=full_text,
                    metadata={"source": path.name}
                ))
                
        except ImportError:
            logger.error("python-docx not installed. Cannot parse DOCX.")
        except Exception as e:
            logger.error(f"Error parsing DOCX {path}: {e}")
            
        return chunks

    def _parse_txt(self, path: Path) -> List[DocumentChunk]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [DocumentChunk(text=text, metadata={"source": path.name})]
